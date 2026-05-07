"""Naitei.ai - FastAPI backend for interview preparation SaaS."""
# Force Railway rebuild: 2026-05-03 14:20 UTC

from __future__ import annotations

import ipaddress
import json
import os
import re
import secrets
import subprocess
import tempfile
import time
import uuid
import zipfile
from abc import ABC, abstractmethod
from base64 import b64encode
from enum import Enum
from io import BytesIO
from typing import Any, AsyncIterator, Generic, Optional, TypeVar

import urllib.request
import httpx

# ── Japanese font for PDF generation ──────────────────────────────────────────
# Pre-register at module load so every PDF call reuses the same font object.
_JP_FONT_NAME: str = "Helvetica"  # default (ASCII-only fallback)
_JP_FONT_BOLD: str = "Helvetica-Bold"

def _init_jp_font() -> None:
    """Try to register a CJK-capable font; silently fall back to Helvetica."""
    global _JP_FONT_NAME, _JP_FONT_BOLD
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.pdfbase.ttfonts import TTFont

        # 1. CID built-in (always available in the reportlab package)
        try:
            pdfmetrics.registerFont(UnicodeCIDFont("HeiseiKakuGo-W5"))
            _JP_FONT_NAME = "HeiseiKakuGo-W5"
            _JP_FONT_BOLD = "HeiseiKakuGo-W5"
            return
        except Exception:
            pass

        # 2. System / cached TTF
        _cache = "/tmp/NotoSansJP-Regular.ttf"
        candidates = [
            _cache,
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        ]
        for path in candidates:
            if os.path.exists(path):
                try:
                    pdfmetrics.registerFont(TTFont("NotoSansJP", path))
                    _JP_FONT_NAME = "NotoSansJP"
                    _JP_FONT_BOLD = "NotoSansJP"
                    return
                except Exception:
                    pass

        # 3. Download NotoSansJP in a background thread (non-blocking)
        # This avoids blocking server startup on Railway
        import threading
        def _download_font():
            global _JP_FONT_NAME, _JP_FONT_BOLD
            try:
                url = "https://raw.githubusercontent.com/google/fonts/main/ofl/notosansjp/static/NotoSansJP-Regular.ttf"
                req2 = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
                with urllib.request.urlopen(req2, timeout=20) as resp:
                    data = resp.read()
                with open(_cache, "wb") as fh:
                    fh.write(data)
                from reportlab.pdfbase import pdfmetrics as _pm
                from reportlab.pdfbase.ttfonts import TTFont as _TTF
                _pm.registerFont(_TTF("NotoSansJP", _cache))
                _JP_FONT_NAME = "NotoSansJP"
                _JP_FONT_BOLD = "NotoSansJP"
            except Exception:
                pass
        threading.Thread(target=_download_font, daemon=True).start()
    except Exception:
        pass  # Helvetica fallback remains

_init_jp_font()
# ──────────────────────────────────────────────────────────────────────────────
from anthropic import AsyncAnthropic
from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException, Request, Response, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel, ConfigDict, Field
from pydantic.alias_generators import to_camel
from docx import Document as DocxDocument
from PyPDF2 import PdfReader

# Load environment variables from .env file
load_dotenv(dotenv_path="/Users/nabehiro/Desktop/Next_Career 2/.env")

# Parse .env file manually as fallback
def load_env_from_file(filepath: str) -> dict[str, str]:
    """Load environment variables from .env file."""
    env_vars = {}
    try:
        with open(filepath, "r") as f:
            for line in f:
                line = line.strip()
                if line and "=" in line and not line.startswith("#"):
                    key, value = line.split("=", 1)
                    env_vars[key.strip()] = value.strip()
    except FileNotFoundError:
        pass
    return env_vars

# Try to load from file first, then from environment
env_file_path = "/Users/nabehiro/Desktop/Next_Career 2/.env"
env_file_vars = load_env_from_file(env_file_path) if os.path.exists(env_file_path) else {}

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------

ANTHROPIC_API_KEY = env_file_vars.get("ANTHROPIC_API_KEY") or os.getenv("ANTHROPIC_API_KEY", "")
OPENAI_API_KEY = env_file_vars.get("OPENAI_API_KEY") or os.getenv("OPENAI_API_KEY", "")
GOOGLE_API_KEY = env_file_vars.get("GOOGLE_GENERATIVE_AI_API_KEY") or os.getenv("GOOGLE_GENERATIVE_AI_API_KEY", "")
RATE_LIMIT_DISABLED = (env_file_vars.get("RATE_LIMIT_DISABLED") or os.getenv("RATE_LIMIT_DISABLED", "")).lower() == "true"
DEFAULT_PROVIDER = env_file_vars.get("AI_PROVIDER") or os.getenv("AI_PROVIDER", "claude")
FALLBACK_ORDER: list[str] = (env_file_vars.get("AI_FALLBACK_ORDER") or os.getenv("AI_FALLBACK_ORDER", "claude,openai,gemini")).split(",")

# ---------------------------------------------------------------------------
# Pydantic Models  (CRITICAL #1: SessionState, CRITICAL #2: InterviewMaterials)
# ---------------------------------------------------------------------------

class CamelModel(BaseModel):
    model_config = ConfigDict(
        alias_generator=to_camel,
        populate_by_name=True,
    )


class CareerEntry(CamelModel):
    company: str
    period_from: str = Field(alias="periodFrom")
    period_to: Optional[str] = Field(None, alias="periodTo")
    role: str
    achievements: list[str] = []


class ApplicantProfile(CamelModel):
    name: str
    age: Optional[int] = Field(None, ge=15, le=99)
    email: Optional[str] = None
    phone: Optional[str] = None
    address: Optional[str] = None
    career_history: list[CareerEntry] = Field(default_factory=list, alias="careerHistory")
    qualifications: list[str] = []
    self_pr: Optional[str] = Field(None, alias="selfPr")
    raw_text: str = Field("", alias="rawText")


class CompanyInfo(CamelModel):
    name: str
    url: Optional[str] = None
    job_title: str = Field("", alias="jobTitle")
    job_description: str = Field("", alias="jobDescription")
    requirements: list[str] = []
    preferred_skills: list[str] = Field(default_factory=list, alias="preferredSkills")
    employment_type: Optional[str] = Field(None, alias="employmentType")
    salary: Optional[str] = None
    selection_flow: list[str] = Field(default_factory=list, alias="selectionFlow")
    raw_text: str = Field("", alias="rawText")
    interview_date: str = Field("", alias="interviewDate")


class MarkdownDoc(CamelModel):
    markdown: str


class InterviewMaterials(CamelModel):
    qa: Optional[MarkdownDoc] = None
    pr: Optional[MarkdownDoc] = None
    questions: Optional[MarkdownDoc] = None
    checklist: Optional[MarkdownDoc] = None


class ReviewComment(CamelModel):
    section: str
    severity: str
    message: str


class ReviewData(CamelModel):
    resume: Optional[str] = None
    career_history: Optional[str] = Field(None, alias="careerHistory")
    comments: list[ReviewComment] = []


class ChatMessage(CamelModel):
    role: str
    content: str
    timestamp: Optional[str] = None


class InterviewState(CamelModel):
    qa: Optional[str] = None
    pr: Optional[str] = None
    questions: Optional[str] = None
    checklist: Optional[str] = None
    chat_history: list[ChatMessage] = Field(default_factory=list, alias="chatHistory")


class Preferences(CamelModel):
    ai_provider: str = Field("claude", alias="aiProvider")
    tier: str = "balanced"


class SessionState(CamelModel):
    version: int = 1
    created_at: str = Field("", alias="createdAt")
    updated_at: str = Field("", alias="updatedAt")
    current_step: str = Field("step1", alias="currentStep")
    applicant: Optional[ApplicantProfile] = None
    company: Optional[CompanyInfo] = None
    review: Optional[ReviewData] = None
    interview: Optional[InterviewState] = None
    preferences: Preferences = Field(default_factory=Preferences)


class UsageInfo(CamelModel):
    prompt_tokens: int = Field(0, alias="promptTokens")
    completion_tokens: int = Field(0, alias="completionTokens")
    total_tokens: int = Field(0, alias="totalTokens")
    cost_usd: Optional[float] = Field(None, alias="costUsd")


class ApiMeta(CamelModel):
    request_id: str = Field("", alias="requestId")
    latency_ms: Optional[float] = Field(None, alias="latencyMs")
    usage: Optional[UsageInfo] = None
    model: Optional[str] = None
    provider: Optional[str] = None


T = TypeVar("T")


class ApiSuccess(CamelModel, Generic[T]):
    success: bool = True
    data: T
    error: None = None
    meta: Optional[ApiMeta] = None


class ApiError(CamelModel):
    success: bool = False
    data: None = None
    error: str
    meta: Optional[ApiMeta] = None


# --- Request models ---

class ScrapeRequest(CamelModel):
    url: str


class ExtractRequest(CamelModel):
    pass


class ReviewRequest(CamelModel):
    applicant: ApplicantProfile
    company: CompanyInfo
    target: str = "both"
    tone: str = "standard"


class InterviewGenRequest(CamelModel):
    applicant: ApplicantProfile
    company: CompanyInfo
    interviewer_profile: Optional[str] = Field(None, alias="interviewerProfile")


class ChatRequest(CamelModel):
    applicant: ApplicantProfile
    company: CompanyInfo
    messages: list[ChatMessage] = []


class PdfRequest(CamelModel):
    materials: InterviewMaterials
    applicant_name: str = Field("", alias="applicantName")
    company_name: str = Field("", alias="companyName")
    interview_date: Optional[str] = Field(None, alias="interviewDate")
    job_title: str = Field("", alias="jobTitle")


class InterviewSlidePdfRequest(BaseModel):
    qa: str = ""
    pr: str = ""
    questions: str = ""
    checklist: str = ""
    applicant_name: str = ""
    company_name: str = ""
    job_title: str = ""
    interview_date: str = ""


class ResearchRequest(CamelModel):
    company: CompanyInfo


# ---------------------------------------------------------------------------
# PII Masker  (HIGH #6)
# ---------------------------------------------------------------------------

class PIIMasker:
    _PATTERNS: list[tuple[str, re.Pattern[str]]] = [
        ("EMAIL", re.compile(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+")),
        ("PHONE", re.compile(r"0\d{1,4}[-\s]?\d{1,4}[-\s]?\d{3,4}")),
        ("POSTAL", re.compile(r"\d{3}-?\d{4}")),
        # BIRTHDAY: Match Japanese date formats (日付記載) but NOT ISO format (YYYY-MM-DD used in APIs)
        # Negative lookahead (?!.*\d-\d) to exclude patterns like "2025-06-01"
        ("BIRTHDAY", re.compile(
            r"(19|20)\d{2}[/年](0?[1-9]|1[0-2])[/月](0?[1-9]|[12]\d|3[01])日?"
        )),
        ("NAME_JP", re.compile(
            r"(?:氏名|名前)\s*[:：]\s*[　-鿿]{1,4}\s*[　-鿿]{1,4}"
        )),
    ]

    @classmethod
    def mask(cls, text: str) -> str:
        for label, pattern in cls._PATTERNS:
            text = pattern.sub(f"[{label}_MASKED]", text)
        return text


# ---------------------------------------------------------------------------
# Token Bucket Rate Limiter  (HIGH #5)
# ---------------------------------------------------------------------------

_RATE_LIMITS: dict[str, tuple[int, int]] = {
    "/api/scrape": (10, 600),
    "/api/extract": (20, 600),
    "/api/review": (10, 600),
    "/api/pdf": (30, 600),
    "/api/interview": (60, 600),
    "/api/research": (5, 600),
}

_buckets: dict[str, dict[str, float]] = {}


def _get_client_ip(request: Request) -> str:
    forwarded = request.headers.get("x-forwarded-for")
    if forwarded:
        return forwarded.split(",")[0].strip()
    return request.client.host if request.client else "unknown"


def _check_rate_limit(ip: str, path: str) -> tuple[bool, int, float]:
    prefix = path
    for key in _RATE_LIMITS:
        if path.startswith(key):
            prefix = key
            break
    else:
        return True, 0, 0.0

    limit, window = _RATE_LIMITS[prefix]
    bucket_key = f"{ip}:{prefix}"
    now = time.monotonic()
    bucket = _buckets.get(bucket_key, {"tokens": float(limit), "last": now})
    elapsed = now - bucket["last"]
    bucket["tokens"] = min(float(limit), bucket["tokens"] + (elapsed / window) * limit)
    bucket["last"] = now

    if bucket["tokens"] < 1.0:
        retry_after = (1.0 - bucket["tokens"]) / (limit / window)
        _buckets[bucket_key] = bucket
        return False, limit, retry_after

    bucket["tokens"] -= 1.0
    _buckets[bucket_key] = bucket
    return True, limit, 0.0


# ---------------------------------------------------------------------------
# SSRF Protection
# ---------------------------------------------------------------------------

_PRIVATE_NETWORKS = [
    ipaddress.ip_network("10.0.0.0/8"),
    ipaddress.ip_network("172.16.0.0/12"),
    ipaddress.ip_network("192.168.0.0/16"),
    ipaddress.ip_network("127.0.0.0/8"),
    ipaddress.ip_network("169.254.0.0/16"),
    ipaddress.ip_network("::1/128"),
]


def _is_safe_url(url: str) -> bool:
    from urllib.parse import urlparse

    parsed = urlparse(url)
    if parsed.scheme != "https":
        return False
    hostname = parsed.hostname
    if not hostname:
        return False
    try:
        import socket
        resolved = socket.getaddrinfo(hostname, None)
        for _, _, _, _, addr in resolved:
            ip = ipaddress.ip_address(addr[0])
            for net in _PRIVATE_NETWORKS:
                if ip in net:
                    return False
    except (socket.gaierror, ValueError):
        return False
    return True


# ---------------------------------------------------------------------------
# Magic Byte File Validation
# ---------------------------------------------------------------------------

MAGIC_BYTES: dict[str, bytes] = {
    "application/pdf": b"%PDF",
    "image/png": b"\x89PNG",
    "image/jpeg": b"\xff\xd8\xff",
}

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB


def validate_upload(content: bytes, content_type: str) -> None:
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(413, "File exceeds 10MB limit")
    expected = MAGIC_BYTES.get(content_type)
    if expected and not content[: len(expected)] == expected:
        raise HTTPException(422, "File signature does not match declared MIME type")


# ---------------------------------------------------------------------------
# File Text Extraction
# ---------------------------------------------------------------------------

def extract_text_from_file(content: bytes, content_type: str) -> str:
    """Extract text from DOCX, PDF, or plain text files."""
    try:
        if content_type in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/x-docx"):
            # Extract text from DOCX
            doc = DocxDocument(BytesIO(content))
            text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
            # Also extract table content
            for table in doc.tables:
                for row in table.rows:
                    text += "\n" + " | ".join(cell.text for cell in row.cells)
            return text.strip()
        elif content_type in ("application/pdf", "text/plain"):
            # Extract text from PDF
            if content_type == "application/pdf":
                pdf_reader = PdfReader(BytesIO(content))
                text = "\n".join(page.extract_text() for page in pdf_reader.pages if page.extract_text())
            else:
                # Plain text
                text = content.decode("utf-8", errors="replace")
            return text.strip()
        else:
            # Default: try to decode as text
            return content.decode("utf-8", errors="replace").strip()
    except Exception as e:
        raise ValueError(f"Failed to extract text from file: {str(e)}")


# ---------------------------------------------------------------------------
# API Response Helpers
# ---------------------------------------------------------------------------

def ok(data: Any, *, request_id: str = "", latency_ms: float = 0, **meta_kw: Any) -> dict:
    meta = {"requestId": request_id, "latencyMs": latency_ms, **meta_kw}
    return {"success": True, "data": data, "error": None, "meta": meta}


def err(message: str, *, status: int = 400, request_id: str = "") -> JSONResponse:
    body = {"success": False, "data": None, "error": PIIMasker.mask(message), "meta": {"requestId": request_id}}
    return JSONResponse(body, status_code=status)


# ---------------------------------------------------------------------------
# AIProvider ABC + Implementations  (Cascade Fallback)
# ---------------------------------------------------------------------------

class TaskKind(str, Enum):
    EXTRACT_JOB = "extract_job_posting"
    EXTRACT_RESUME = "extract_resume"
    REVIEW = "review_document"
    GEN_QA = "generate_qa"
    GEN_PR = "generate_self_pr"
    GEN_QUESTIONS = "generate_counter_questions"
    GEN_CHECKLIST = "generate_checklist"
    COMPANY_RESEARCH = "company_research"
    CHAT = "chat"


MODEL_MAP: dict[str, dict[TaskKind, str]] = {
    "claude": {
        TaskKind.EXTRACT_JOB: "claude-haiku-4-5",
        TaskKind.EXTRACT_RESUME: "claude-haiku-4-5",
        TaskKind.REVIEW: "claude-haiku-4-5",
        TaskKind.GEN_QA: "claude-haiku-4-5",
        TaskKind.GEN_PR: "claude-haiku-4-5",
        TaskKind.GEN_QUESTIONS: "claude-haiku-4-5",
        TaskKind.GEN_CHECKLIST: "claude-haiku-4-5",
        TaskKind.COMPANY_RESEARCH: "claude-haiku-4-5",
        TaskKind.CHAT: "claude-haiku-4-5",
    },
    "openai": {k: "gpt-4o" for k in TaskKind},
    "gemini": {k: "gemini-2.5-flash" for k in TaskKind},
}


class AIProvider(ABC):
    name: str

    @abstractmethod
    async def generate_text(self, prompt: str, *, system: str = "", model: str = "", **kw: Any) -> dict:
        ...

    @abstractmethod
    async def stream_text(self, prompt: str, *, system: str = "", model: str = "", **kw: Any) -> AsyncIterator[dict]:
        ...

    @abstractmethod
    async def generate_from_image(self, image: bytes, mime: str, prompt: str, *, model: str = "", **kw: Any) -> dict:
        ...


class ClaudeProvider(AIProvider):
    name = "claude"

    async def generate_text(self, prompt: str, *, system: str = "", model: str = "claude-haiku-4-5", **kw: Any) -> dict:
        # Adjust max_tokens based on model
        max_tokens = 2048 if "haiku" in model else 4096

        # Create explicit http_client to avoid deprecated proxies parameter issue
        http_client = httpx.AsyncClient(timeout=120)
        try:
            client = AsyncAnthropic(api_key=ANTHROPIC_API_KEY, http_client=http_client)
            print(f"[DEBUG] AsyncAnthropic client created successfully")
        except Exception as e:
            print(f"[ERROR] Failed to create AsyncAnthropic client: {e}")
            await http_client.aclose()
            raise

        try:
            # Build system messages if provided
            system_blocks = None
            if system:
                system_blocks = [{"type": "text", "text": system, "cache_control": {"type": "ephemeral"}}]

            # Call Anthropic API using SDK
            response = await client.messages.create(
                model=model,
                max_tokens=max_tokens,
                system=system_blocks if system_blocks else None,
                messages=[{"role": "user", "content": prompt}]
            )

            # Extract text from response
            text = "".join(b.text for b in response.content if b.type == "text")
            return {"text": text, "usage": {"input_tokens": response.usage.input_tokens, "output_tokens": response.usage.output_tokens}}
        finally:
            await http_client.aclose()

    async def stream_text(self, prompt: str, *, system: str = "", model: str = "claude-haiku-4-5", **kw: Any) -> AsyncIterator[dict]:
        # Adjust max_tokens based on model
        max_tokens = 2048 if "haiku" in model else 4096

        # Create explicit http_client to avoid deprecated proxies parameter issue
        http_client = httpx.AsyncClient(timeout=120)
        try:
            client = AsyncAnthropic(api_key=ANTHROPIC_API_KEY, http_client=http_client)
            print(f"[DEBUG] AsyncAnthropic client created successfully")
        except Exception as e:
            print(f"[ERROR] Failed to create AsyncAnthropic client: {e}")
            await http_client.aclose()
            raise

        print(f"[DEBUG ClaudeProvider.stream_text] Sending request to Anthropic API:")
        print(f"  Model: {model}")
        print(f"  API Key: {ANTHROPIC_API_KEY[:30]}...{ANTHROPIC_API_KEY[-20:]}")
        print(f"  System: {system[:100] if system else 'None'}...")
        print(f"  Prompt length: {len(prompt)}")

        # Build system messages if provided
        system_param = None
        if system:
            system_param = [{"type": "text", "text": system, "cache_control": {"type": "ephemeral"}}]

        # Use SDK streaming with async context manager - automatically handles version negotiation
        try:
            print(f"[DEBUG] Starting messages.stream with model={model}, max_tokens={max_tokens}")
            async with client.messages.stream(
                model=model,
                max_tokens=max_tokens,
                system=system_param if system_param else None,
                messages=[{"role": "user", "content": prompt}]
            ) as stream:
                print(f"[DEBUG] Stream context manager entered successfully")
                # Stream text deltas
                async for text in stream.text_stream:
                    if text:
                        yield {"type": "text-delta", "delta": text}
                print(f"[DEBUG] Stream text_stream iteration completed")
                # get_final_message は with ブロック内で呼ぶ（ブロック外では無効になる場合がある）
                try:
                    final_message = await stream.get_final_message()
                    yield {"type": "finish", "usage": {"input_tokens": final_message.usage.input_tokens, "output_tokens": final_message.usage.output_tokens}}
                except Exception:
                    yield {"type": "finish", "usage": {}}
        except Exception as e:
            print(f"[ERROR] Exception during streaming: {type(e).__name__}: {e}")
            import traceback
            traceback.print_exc()
            raise
        finally:
            await http_client.aclose()

    async def generate_from_image(self, image: bytes, mime: str, prompt: str, *, model: str = "claude-haiku-4-5", **kw: Any) -> dict:
        # Adjust max_tokens based on model
        max_tokens = 2048 if "haiku" in model else 4096
        http_client = httpx.AsyncClient(timeout=120)
        try:
            client = AsyncAnthropic(api_key=ANTHROPIC_API_KEY, http_client=http_client)

            # Encode image to base64
            b64 = b64encode(image).decode()

            # Call Anthropic API using SDK with image
            response = await client.messages.create(
                model=model,
                max_tokens=max_tokens,
                messages=[{
                    "role": "user",
                    "content": [
                        {"type": "image", "source": {"type": "base64", "media_type": mime, "data": b64}},
                        {"type": "text", "text": prompt},
                    ]
                }]
            )

            # Extract text from response
            text = "".join(b.text for b in response.content if b.type == "text")
            return {"text": text, "usage": {"input_tokens": response.usage.input_tokens, "output_tokens": response.usage.output_tokens}}
        finally:
            await http_client.aclose()


class OpenAIProvider(AIProvider):
    name = "openai"

    async def _chat(self, messages: list[dict], *, model: str, stream: bool = False, **kw: Any) -> httpx.Response:
        async with httpx.AsyncClient(timeout=120) as c:
            return await c.post("https://api.openai.com/v1/chat/completions", json={
                "model": model, "messages": messages, "stream": stream, **kw,
            }, headers={"Authorization": f"Bearer {OPENAI_API_KEY}", "Content-Type": "application/json"})

    async def generate_text(self, prompt: str, *, system: str = "", model: str = "gpt-4o", **kw: Any) -> dict:
        msgs: list[dict] = []
        if system:
            msgs.append({"role": "system", "content": system})
        msgs.append({"role": "user", "content": prompt})
        r = await self._chat(msgs, model=model)
        r.raise_for_status()
        data = r.json()
        return {"text": data["choices"][0]["message"]["content"], "usage": data.get("usage", {})}

    async def stream_text(self, prompt: str, *, system: str = "", model: str = "gpt-4o", **kw: Any) -> AsyncIterator[dict]:
        msgs: list[dict] = []
        if system:
            msgs.append({"role": "system", "content": system})
        msgs.append({"role": "user", "content": prompt})
        async with httpx.AsyncClient(timeout=120) as c:
            async with c.stream("POST", "https://api.openai.com/v1/chat/completions", json={
                "model": model, "messages": msgs, "stream": True,
            }, headers={"Authorization": f"Bearer {OPENAI_API_KEY}", "Content-Type": "application/json"}) as resp:
                resp.raise_for_status()
                async for line in resp.aiter_lines():
                    if not line.startswith("data: ") or line.strip() == "data: [DONE]":
                        continue
                    chunk = json.loads(line[6:])
                    delta = chunk.get("choices", [{}])[0].get("delta", {}).get("content", "")
                    if delta:
                        yield {"type": "text-delta", "delta": delta}
        yield {"type": "finish", "usage": {}}

    async def generate_from_image(self, image: bytes, mime: str, prompt: str, *, model: str = "gpt-4o", **kw: Any) -> dict:
        b64 = b64encode(image).decode()
        msgs = [{"role": "user", "content": [
            {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
            {"type": "text", "text": prompt},
        ]}]
        r = await self._chat(msgs, model=model)
        r.raise_for_status()
        data = r.json()
        return {"text": data["choices"][0]["message"]["content"], "usage": data.get("usage", {})}


class GeminiProvider(AIProvider):
    name = "gemini"

    async def generate_text(self, prompt: str, *, system: str = "", model: str = "gemini-2.5-flash", **kw: Any) -> dict:
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={GOOGLE_API_KEY}"
        parts: list[dict] = []
        if system:
            parts.append({"text": system})
        parts.append({"text": prompt})
        async with httpx.AsyncClient(timeout=120) as c:
            r = await c.post(url, json={"contents": [{"parts": parts}]})
            r.raise_for_status()
            data = r.json()
            text = data.get("candidates", [{}])[0].get("content", {}).get("parts", [{}])[0].get("text", "")
            return {"text": text, "usage": data.get("usageMetadata", {})}

    async def stream_text(self, prompt: str, *, system: str = "", model: str = "gemini-2.5-flash", **kw: Any) -> AsyncIterator[dict]:
        result = await self.generate_text(prompt, system=system, model=model)
        yield {"type": "text-delta", "delta": result["text"]}
        yield {"type": "finish", "usage": result.get("usage", {})}

    async def generate_from_image(self, image: bytes, mime: str, prompt: str, *, model: str = "gemini-2.5-flash", **kw: Any) -> dict:
        b64 = b64encode(image).decode()
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={GOOGLE_API_KEY}"
        parts = [{"inline_data": {"mime_type": mime, "data": b64}}, {"text": prompt}]
        async with httpx.AsyncClient(timeout=120) as c:
            r = await c.post(url, json={"contents": [{"parts": parts}]})
            r.raise_for_status()
            data = r.json()
            text = data.get("candidates", [{}])[0].get("content", {}).get("parts", [{}])[0].get("text", "")
            return {"text": text, "usage": data.get("usageMetadata", {})}


_PROVIDERS: dict[str, AIProvider] = {
    "claude": ClaudeProvider(),
    "openai": OpenAIProvider(),
    "gemini": GeminiProvider(),
}


def get_provider(name: str | None = None) -> AIProvider:
    return _PROVIDERS[name or DEFAULT_PROVIDER]


async def with_fallback(fn, *, skip: list[str] | None = None):
    last_err: Exception | None = None
    for name in FALLBACK_ORDER:
        if skip and name in skip:
            continue
        try:
            return await fn(_PROVIDERS[name])
        except (httpx.HTTPStatusError, httpx.ConnectError) as e:
            last_err = e
            status = getattr(e, "response", None)
            if status and status.status_code not in (429, 500, 502, 503):
                raise
    raise last_err or RuntimeError("No AI providers available")


# ---------------------------------------------------------------------------
# SSE Streaming Helper
# ---------------------------------------------------------------------------

async def sse_stream(provider: AIProvider, prompt: str, system: str, task: TaskKind, request_id: str):
    model = MODEL_MAP.get(provider.name, {}).get(task, "")

    async def generate():
        try:
            async for chunk in provider.stream_text(prompt, system=system, model=model):
                if chunk["type"] == "text-delta":
                    masked = PIIMasker.mask(chunk["delta"])
                    yield f"event: content\ndata: {json.dumps({'type': 'content', 'delta': masked})}\n\n"
                elif chunk["type"] == "finish":
                    yield f"event: done\ndata: {json.dumps({'type': 'done', 'usage': chunk.get('usage', {}), 'requestId': request_id})}\n\n"
        except Exception as exc:
            yield f"event: error\ndata: {json.dumps({'type': 'error', 'message': PIIMasker.mask(str(exc))})}\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream", headers={
        "Cache-Control": "no-cache", "X-Request-Id": request_id,
    })


# ---------------------------------------------------------------------------
# FastAPI App
# ---------------------------------------------------------------------------

app = FastAPI(title="Naitei.ai", version="0.1.0")

app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])


# Global exception handler — ensures all unhandled exceptions return JSON (not plain text)
@app.exception_handler(Exception)
async def global_exception_handler(request: Request, exc: Exception) -> JSONResponse:
    rid = getattr(request.state, "request_id", "")
    return JSONResponse(
        {"success": False, "data": None, "error": f"Internal server error: {type(exc).__name__}", "meta": {"requestId": rid}},
        status_code=500,
    )


# --- CSP Nonce + Security Headers Middleware (CRITICAL #4) ---

@app.middleware("http")
async def security_middleware(request: Request, call_next):
    request_id = str(uuid.uuid4())
    request.state.request_id = request_id
    nonce = secrets.token_urlsafe(16)
    request.state.csp_nonce = nonce
    start = time.monotonic()

    if not RATE_LIMIT_DISABLED and request.url.path.startswith("/api/"):
        ip = _get_client_ip(request)
        allowed, limit, retry_after = _check_rate_limit(ip, request.url.path)
        if not allowed:
            return JSONResponse(
                {"success": False, "data": None, "error": "Rate limit exceeded", "meta": {"requestId": request_id}},
                status_code=429,
                headers={"Retry-After": str(int(retry_after) + 1), "X-RateLimit-Limit": str(limit)},
            )

    response = await call_next(request)
    elapsed = (time.monotonic() - start) * 1000

    if not request.url.path.startswith(("/docs", "/redoc", "/openapi.json")):
        response.headers["Content-Security-Policy"] = (
            f"default-src 'self'; "
            f"script-src 'nonce-{nonce}' 'strict-dynamic'; "
            f"style-src 'self' 'nonce-{nonce}'; "
            f"img-src 'self' data: https:; "
            f"font-src 'self' data:; "
            f"connect-src 'self' https://api.anthropic.com https://api.openai.com https://generativelanguage.googleapis.com; "
            f"frame-ancestors 'none'; object-src 'none'; base-uri 'self'"
        )
    response.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains; preload"
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["X-Frame-Options"] = "DENY"
    response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
    response.headers["Permissions-Policy"] = "camera=(), microphone=(), geolocation=()"
    response.headers["X-Request-Id"] = request_id
    response.headers["X-Response-Time"] = f"{elapsed:.0f}ms"
    return response


# ---------------------------------------------------------------------------
# Endpoints
# ---------------------------------------------------------------------------

@app.get("/api/health")
async def health():
    return {"status": "ok", "jp_font": _JP_FONT_NAME}


@app.post("/api/scrape")
async def scrape(body: ScrapeRequest, request: Request):
    rid = request.state.request_id
    if not _is_safe_url(body.url):
        return err("URL is not allowed (must be HTTPS, non-private)", status=422, request_id=rid)
    start = time.monotonic()
    async with httpx.AsyncClient(timeout=30, follow_redirects=True) as c:
        try:
            r = await c.get(body.url, headers={"User-Agent": "Naitei.ai/0.1"})
            r.raise_for_status()
        except httpx.HTTPError as e:
            return err(f"Failed to fetch URL: {e}", status=502, request_id=rid)
    provider = get_provider()
    model = MODEL_MAP.get(provider.name, {}).get(TaskKind.EXTRACT_JOB, "")
    result = await provider.generate_text(
        f"Extract structured job posting information from the following HTML/text. Return JSON with fields: jobTitle, company, requirements, preferredSkills, employmentType, salary, jobDescription.\n\n{r.text[:15000]}",
        system="You are a job posting extraction specialist. Return valid JSON only.",
        model=model,
    )
    elapsed = (time.monotonic() - start) * 1000
    return ok({"raw_text": r.text[:5000], "structured": result["text"]}, request_id=rid, latency_ms=elapsed, provider=provider.name, model=model)


@app.post("/api/extract")
async def extract(request: Request, file: UploadFile):
    rid = request.state.request_id
    content = await file.read()
    validate_upload(content, file.content_type or "application/octet-stream")

    # Extract text from file (DOCX, PDF, or plain text)
    try:
        file_text = extract_text_from_file(content, file.content_type or "text/plain")
    except ValueError as e:
        return err(f"File extraction failed: {str(e)}", status=422, request_id=rid)

    if not file_text:
        return err("Uploaded file appears to be empty", status=422, request_id=rid)

    provider = get_provider()
    model = MODEL_MAP.get(provider.name, {}).get(TaskKind.EXTRACT_RESUME, "")

    # Use generate_text() with extracted file content
    prompt = f"""Extract applicant profile from the following document content.
Return ONLY a JSON object (no markdown, no code blocks) with these exact fields:
- name (string, required)
- age (number or null)
- careerHistory (array of objects with: company, periodFrom, periodTo, role, achievements)
- qualifications (array of strings or null)
- selfPr (string or null)

Document content:
{file_text[:4000]}"""  # Limit to 4000 chars to avoid token overrun

    try:
        result = await provider.generate_text(prompt, model=model)
    except Exception as e:
        # AI failed but we still have the raw file text — return a minimal profile so step3 can proceed
        fallback = ApplicantProfile(name="", raw_text=file_text[:4000])
        content_type = file.content_type or "text/plain"
        return ok({
            "applicant": fallback.model_dump(by_alias=True),
            "sources": [{"fileName": file.filename or "uploaded-file", "mimeType": content_type, "pageCount": 1, "extractionMethod": "fallback-raw"}],
            "warnings": [f"AI extraction failed ({type(e).__name__}), raw text preserved for review"]
        }, request_id=rid, provider=provider.name, model=model)

    # Parse LLM output as JSON and validate structure
    llm_text = result["text"].strip()
    applicant_data = None

    try:
        # Try to extract JSON from LLM output (may be wrapped in markdown code blocks)
        json_match = re.search(r'```(?:json)?\s*([\s\S]*?)```', llm_text)
        if json_match:
            json_str = json_match.group(1).strip()
        else:
            json_str = llm_text

        parsed = json.loads(json_str)

        # Validate against ApplicantProfile schema
        # Inject rawText from actual file content so frontend can use it for review
        parsed["rawText"] = file_text[:4000]
        applicant_data = ApplicantProfile(**parsed)
    except (json.JSONDecodeError, ValueError) as e:
        # If JSON parsing fails, return error with raw text for debugging
        return err(f"Failed to parse extracted data as JSON: {str(e)}", status=422, request_id=rid)
    except Exception as e:
        return err(f"Failed to validate extracted data: {str(e)}", status=422, request_id=rid)

    # Determine extraction method based on file type
    content_type = file.content_type or "text/plain"
    if "pdf" in content_type.lower():
        extraction_method = "pdf-text"
    elif "wordprocessingml" in content_type.lower() or "docx" in content_type.lower():
        extraction_method = "docx-text"
    else:
        extraction_method = "text"

    # Return success with complete ExtractResponse structure
    return ok({
        "applicant": applicant_data.model_dump(by_alias=True),
        "sources": [{
            "fileName": file.filename or "uploaded-file",
            "mimeType": content_type,
            "pageCount": 1,
            "extractionMethod": extraction_method
        }],
        "warnings": []
    }, request_id=rid, provider=provider.name, model=model)


@app.post("/api/review")
async def review(body: ReviewRequest, request: Request):
    rid = request.state.request_id
    provider = get_provider()
    applicant = body.applicant
    company = body.company
    # Use rawText as fallback if structured fields are empty
    applicant_text = (
        applicant.raw_text
        or applicant.self_pr
        or (", ".join(f"{c.role} at {c.company}" for c in applicant.career_history) if applicant.career_history else "")
        or "(応募者情報なし)"
    )
    company_text = (
        company.raw_text
        or company.job_description
        or f"{company.name} {company.job_title}".strip()
        or "(求人情報なし)"
    )
    prompt = (
        f"【応募者情報】\n{applicant_text[:3000]}\n\n"
        f"【求人情報】\n{company_text[:2000]}\n\n"
        f"ターゲット: {body.target}, トーン: {body.tone}\n\n"
        "上記の応募者の職務経歴書・志望動機を、求人要件に合わせて日本語で添削してください。"
        "改善後のテキスト、具体的な変更点、コメント（重要度別）、適合度スコア（0-100）、総評を提供してください。"
        "応募者情報や求人情報が不完全でも、提供された情報をもとに最善の添削を行ってください。"
    )
    return await sse_stream(provider, prompt, "あなたは日本の就職・転職支援の専門家です。応募者の書類を丁寧に添削します。", TaskKind.REVIEW, rid)


@app.post("/api/interview/qa")
async def interview_qa(body: InterviewGenRequest, request: Request):
    rid = request.state.request_id
    provider = get_provider()
    applicant_json = body.applicant.model_dump_json(by_alias=True)
    company_json = body.company.model_dump_json(by_alias=True)
    interviewer = body.interviewer_profile or "未判明（役職・部署・組織ミッションから想定プロファイルを構築すること）"
    applicant_name = body.applicant.name or "応募者"
    company_name = body.company.name or "応募先"
    job_title = body.company.job_title or "応募職種"
    interview_date = body.company.interview_date or ""
    prompt = f"""応募者情報: {applicant_json}
応募先情報: {company_json}
面接官: {interviewer}

⚠️ 出力ルール（必ず守ること）:
- コードブロック（```markdown や ``` など）で囲まない
- 前置き文・説明文・コメントを一切書かない
- 最初の文字からMarpスライドの内容を直接出力する
- 出力の最初は `<!-- _class: title -->` で始まる

# 出力形式：Marpスライド形式（想定問答集）

スライド区切り文字は `---`（前後に空行）。目標ページ数：**25〜35スライド**。
スライドCSSクラスの指定は `<!-- _class: クラス名 -->` をスライド冒頭に記述。

使用できるCSSクラス：
- `<!-- _class: title -->` … 表紙（グラデーション紺背景）
- `<div class="tip">` … 黄背景ヒントボックス
- `<div class="sep"></div>` … 破線セパレーター（同一スライドで2問ペア時に使用）
- `<div class="two-col"><div>左</div><div>右</div></div>` … 2カラム
- `<div class="magnet">` … 磁力ワード（黄破線枠）
- blockquote (`>`) … 面接回答本文

## 構成テンプレート

スライド1: 表紙
```
<!-- _class: title -->

# 想定問答集

## {company_name} 面接対策

**{applicant_name}**　{f'面接日：{interview_date}' if interview_date else ''}
{job_title}
```

スライド2: 5原則 ＆ 磁力キーワード

スライド3〜4: Q早見表（カテゴリ×推奨度×ページ番号）テーブル×2枚

スライド5〜30以降: Q&Aスライド群
- カテゴリ見出し（A〜H）ごとにh2見出しを入れる
- 4行以下の回答は2問1スライドにペアリング（`<div class="sep"></div>`で区切る）
- 長い回答は1問1スライド、blockquoteで本文、`<div class="tip">コツ</div>`で補足

最終スライド: リスク早見表 + カバレッジチェック表（テーブル）

## Q&Aスライドの形式例（2問ペア）

```
## A. 導入・定番質問

### Q1. 簡単に自己紹介をお願いします

> 「回答本文を口語体・丁寧体で記述」

<div class="tip">💡 コツ：ポイントを記述</div>

<div class="sep"></div>

### Q2. 次の質問文

> 「回答」

<div class="tip">💡 コツ</div>
```

## 生成ルール
- カテゴリA〜H、最低30問（応募者固有エピソード・職歴・資格を具体的に織り込む）
- Fカテゴリ（リスク質問）は必ず含め、応募者のリスク要因に正面から対処
- 業界・職種固有の観点を追加（公務員：中立公正・法令遵守等）
- 回答は30〜90秒の口語（書面の文章ではなく話し言葉）
- `---` で区切られた各スライドは960×540px相当に収まる量（多すぎない）"""
    return await sse_stream(provider, prompt, "あなたは日本の就職・転職支援の専門家です。応募者固有のエピソードと業界知識を活かした実践的な面接対策スライドを作成します。", TaskKind.GEN_QA, rid)


@app.post("/api/interview/pr")
async def interview_pr(body: InterviewGenRequest, request: Request):
    rid = request.state.request_id
    provider = get_provider()
    applicant_json = body.applicant.model_dump_json(by_alias=True)
    company_json = body.company.model_dump_json(by_alias=True)
    applicant_name = body.applicant.name or "応募者"
    company_name = body.company.name or "応募先"
    job_title = body.company.job_title or "応募職種"
    interview_date = body.company.interview_date or ""
    prompt = f"""応募者情報: {applicant_json}
応募先情報: {company_json}

⚠️ 出力ルール（必ず守ること）:
- コードブロック（```markdown や ``` など）で囲まない
- 前置き文・説明文・コメントを一切書かない
- 最初の文字からMarpスライドの内容を直接出力する
- 出力の最初は `<!-- _class: title -->` で始まる

# 出力形式：Marpスライド形式（自己PR案）

スライド区切り文字は `---`（前後に空行）。目標ページ数：**14〜18スライド**。

使用できるCSSクラス：
- `<!-- _class: title -->` … 表紙
- `<!-- _class: core -->` … コアフレーズ強調（ダークネイビー背景）
  - 内部に `<div class="bigbox">コアフレーズ</div>` を使う
- `<div class="tip">` … ヒントボックス（黄）
- `<div class="ng">` … NG表現ボックス（赤）
- `<div class="ok">` … OKパターンボックス（緑）
- `<div class="magnet">` … 磁力ワード（黄破線枠・大フォント）
- `<div class="two-col"><div>左</div><div>右</div></div>` … 2カラム

## 構成テンプレート（14〜18スライド）

1. 表紙（`<!-- _class: title -->`）
2. コアフレーズ（`<!-- _class: core -->`）: 強みを1フレーズに凝縮、bigboxに入れる
3. 3バージョン使い分け表（テーブル）
4. 60秒版 本文（blockquoteで本文、tip でポイント）
5. 90秒版 本文①（前半）
6. 90秒版 本文②（後半）＋ポイント
7. 3分版 導入＋①前半
8. 3分版 ①エピソード＋②
9. 3分版 ③＋締め
10. 求人要件対応表（テーブル）
11. 音読練習ポイント
12. NG表現リスト（`<div class="ng">`）
13. 当日フロー（テーブル or リスト）
14. 磁力ワード（`<div class="magnet">`）

## 表紙スライド例

```
<!-- _class: title -->

# 自己PR案

## {company_name} 面接対策

**{applicant_name}**　{f'面接日：{interview_date}' if interview_date else ''}
{job_title}
```

## coreスライド例

```
<!-- _class: core -->

# 🗝️ 自己PRの核（冒頭と末尾で必ず言う）

<div class="bigbox">

（強みを1フレーズに凝縮した文章）

</div>
```

## 生成ルール
- 各バージョンの冒頭と末尾に同一の強みフレーズを配置
- 強みは1フレーズに凝縮（例：「現場に寄り添い、採用課題を解決する提案力」）
- 口語体（丁寧体）、話し言葉
- 応募者の職歴・資格・具体エピソードを盛り込む
- 応募先の求人要件・ミッションとの接点を明示
- 各スライドは960×540px相当に収まる量"""
    return await sse_stream(provider, prompt, "あなたは日本の就職・転職支援の専門家です。応募者の強みを最大化する自己PR案スライドを作成します。", TaskKind.GEN_PR, rid)


@app.post("/api/interview/questions")
async def interview_questions(body: InterviewGenRequest, request: Request):
    rid = request.state.request_id
    provider = get_provider()
    applicant_json = body.applicant.model_dump_json(by_alias=True)
    company_json = body.company.model_dump_json(by_alias=True)
    applicant_name = body.applicant.name or "応募者"
    company_name = body.company.name or "応募先"
    job_title = body.company.job_title or "応募職種"
    interview_date = body.company.interview_date or ""
    prompt = f"""応募者情報: {applicant_json}
応募先情報: {company_json}

⚠️ 出力ルール（必ず守ること）:
- コードブロック（```markdown や ``` など）で囲まない
- 前置き文・説明文・コメントを一切書かない
- 最初の文字からMarpスライドの内容を直接出力する
- 出力の最初は `<!-- _class: title -->` で始まる

# 出力形式：Marpスライド形式（逆質問集）

スライド区切り文字は `---`（前後に空行）。目標ページ数：**14〜18スライド**。

使用できるCSSクラス：
- `<!-- _class: title -->` … 表紙
- `<!-- _class: hero -->` … 重要情報ハイライト（紺グラデ背景）
  - 内部に `<div class="star-box">⭐ 内容</div>` を使う
- `<div class="q-card">` … 質問カード（水色枠）
  - q-card内部の見出しは `### Q番号.（★★★）タイトル`
  - q-card内部でblockquoteを質問文に使う
- `<div class="ng">` … NG表現（赤ボーダー）
- `<div class="flow">` … フロー説明（水色枠）
- `<div class="two-col">` … 2カラム

## 構成テンプレート（14〜18スライド）

1. 表紙（`<!-- _class: title -->`）
2. ヒーロー（`<!-- _class: hero -->`）: 鉄板セット★★★と絶対NG
3. 基本方針表（テーブル）
4. 10問一覧表（テーブル）
5〜7. 業務理解質問群（`<div class="q-card">`）× 2〜3スライド
8〜10. 貢献意欲・組織理解・締め質問群 × 2〜3スライド
11. NG逆質問リスト（`<div class="ng">`）
12. 時間別使い分け戦略（フロー＋テーブル）
13. 話し方テンプレート
14. 面接官タイプ別反応予測＋差がつくパターン
15. 最終確認リスト

## 表紙スライド例

```
<!-- _class: title -->

# 逆質問集

## {company_name} 面接対策

**{applicant_name}**　{f'面接日：{interview_date}' if interview_date else ''}
{job_title}
```

## q-cardスライド例

```
<div class="q-card">

### ⭐ Q1.【★★★最推奨】質問タイトル

> 「質問文をここに書く」

**狙い**：面接官に与えたい印象・引き出したい情報。

</div>

<div class="q-card">

### Q2.【★★☆】質問タイトル

> 「質問文」

**狙い**：...

</div>
```

## 生成ルール
- 応募先の具体的な事業・施策・職種名を質問に織り込む
- ★★★は最低3問用意
- ★評価基準：★★★=最推奨（必ず聞く）/ ★★☆=時間があれば / ★☆☆=状況次第
- 各スライドは960×540px相当に収まる量"""
    return await sse_stream(provider, prompt, "あなたは日本の就職・転職支援の専門家です。面接官の印象を高める逆質問集スライドを作成します。", TaskKind.GEN_QUESTIONS, rid)


@app.post("/api/interview/checklist")
async def interview_checklist(body: InterviewGenRequest, request: Request):
    rid = request.state.request_id
    provider = get_provider()
    applicant_json = body.applicant.model_dump_json(by_alias=True)
    company_json = body.company.model_dump_json(by_alias=True)
    applicant_name = body.applicant.name or "応募者"
    company_name = body.company.name or "応募先"
    job_title = body.company.job_title or "応募職種"
    interview_date = body.company.interview_date or ""
    prompt = f"""応募者情報: {applicant_json}
応募先情報: {company_json}

⚠️ 出力ルール（必ず守ること）:
- コードブロック（```markdown や ``` など）で囲まない
- 前置き文・説明文・コメントを一切書かない
- 最初の文字からMarpスライドの内容を直接出力する
- 出力の最初は `<!-- _class: title -->` で始まる

# 出力形式：Marpスライド形式（事前準備チェックリスト）

スライド区切り文字は `---`（前後に空行）。目標ページ数：**16〜22スライド**。

使用できるCSSクラス：
- `<!-- _class: title -->` … 表紙
- `<!-- _class: hero -->` … 重要情報ハイライト（紺グラデ背景）
- `<!-- _class: mantra -->` … 面接直前の心の呪文（末尾専用）
  - 内部に `<div class="mantra-box">① ...\n② ...\n③ ...</div>` を使う
- `<div class="flow">` … 入退室フロー（水色枠）
- `<div class="two-col">` … 2カラム
- `<div class="tip">` … ヒント・注意
- `<div class="ng">` … NG行動（赤ボーダー）
- チェックボックスリスト `- [ ] 項目`

## 構成テンプレート（16〜22スライド）

1. 表紙（`<!-- _class: title -->`）
2. ヒーロー（`<!-- _class: hero -->`）: 面接基本情報・アクセス表
3. タイムライン早見表（テーブル）
4. 3日前: 書類・知識（`<div class="two-col">`）
5. 3日前: 服装・健康
6. 前日: 経路・パッキング（`<div class="two-col">`）
7. 当日朝
8. 入室の流れ（`<div class="flow">`）
9. 着席後の姿勢・話し方（`<div class="two-col">`）＋NG
10. 面接中の応急対応
11. 退室の流れ（`<div class="flow">`）
12. 面接後・合否連絡（`<div class="two-col">`）
13. 緊急時対応（遅刻・体調不良・想定外質問）
14. 前日夜最終チェック＋第一印象
15. 心の呪文（`<!-- _class: mantra -->`）

## 表紙スライド例

```
<!-- _class: title -->

# 事前準備チェックリスト

## {company_name} 面接対策

**{applicant_name}**　{f'面接日：{interview_date}' if interview_date else ''}
{job_title}
```

## flowスライド例（入室の流れ）

```
# 📋 入室の流れ

<div class="flow">

① ドアの前で一呼吸
↓　② ノック3回（3回）
↓　③ 「どうぞ」の声を待つ
↓　④ 「失礼いたします」と入室
↓　⑤ ドアを静かに閉める（後ろ手NG）
↓　⑥ 「よろしくお願いいたします」と一礼

</div>
```

## 生成ルール
- 応募者固有の持ち物（資格証明書・ポートフォリオ等）を明記
- 応募先の業界特性・職種特性に合わせた具体的な内容
- チェックボックス `- [ ]` 形式を徹底
- 各スライドは960×540px相当に収まる量
- 最終スライドは必ず mantra クラスで締める"""
    return await sse_stream(provider, prompt, "あなたは日本の就職・転職支援の専門家です。応募者が万全の準備で面接に臨めるチェックリストスライドを作成します。", TaskKind.GEN_CHECKLIST, rid)


@app.post("/api/interview/chat")
async def interview_chat(body: ChatRequest, request: Request):
    rid = request.state.request_id
    provider = get_provider()
    history = "\n".join(f"[{m.role}]: {m.content}" for m in body.messages)
    prompt = (
        f"応募者情報: {body.applicant.model_dump_json(by_alias=True)}\n"
        f"応募先情報: {body.company.model_dump_json(by_alias=True)}\n"
        f"会話履歴:\n{history}\n"
        "会話を継続し、応募者が面接回答をより深められるようコーチングしてください。"
        "具体的な改善点・別の言い回し・追加エピソードの引き出し方を提案してください。"
    )
    return await sse_stream(provider, prompt, "あなたは日本の就職・転職支援の専門家です。対話形式で面接回答をブラッシュアップするコーチングを行います。", TaskKind.CHAT, rid)


@app.post("/api/research")
async def research(body: ResearchRequest, request: Request):
    rid = request.state.request_id
    provider = get_provider()
    prompt = (
        f"Company: {body.company.model_dump_json(by_alias=True)}\n"
        "Research this company thoroughly: mission/vision/values, recent news, industry trends, "
        "competitors, work culture, interview tendencies, and ideal candidate profile."
    )
    return await sse_stream(provider, prompt, "You are a thorough company research analyst for Japanese job market.", TaskKind.COMPANY_RESEARCH, rid)


# --- /api/pdf  (ReportLab ベース PDF 生成) ---

def _markdown_to_pdf_bytes(title: str, markdown_text: str, header: str, footer: str) -> bytes:
    """マークダウンテキストを ReportLab で PDF バイト列に変換する。Interview_Workflow Marp デザイン準拠。"""
    import re
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle, KeepTogether

    font_name = _JP_FONT_NAME
    bold_font = _JP_FONT_BOLD

    buf = BytesIO()
    page_w = A4[0] - 40 * mm  # usable width
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm,
        topMargin=20 * mm, bottomMargin=20 * mm,
    )

    styles = getSampleStyleSheet()
    sn = ParagraphStyle("jp_n", parent=styles["Normal"], fontName=font_name, fontSize=10, leading=16, spaceAfter=3)
    sh1 = ParagraphStyle("jp_h1", parent=styles["Heading1"], fontName=bold_font, fontSize=16, leading=22,
                         spaceBefore=10, spaceAfter=4, textColor=colors.HexColor("#0b1e4d"))
    sh2 = ParagraphStyle("jp_h2", parent=styles["Heading2"], fontName=bold_font, fontSize=13, leading=18,
                         spaceBefore=8, spaceAfter=3, textColor=colors.HexColor("#1e40af"))
    sh3 = ParagraphStyle("jp_h3", parent=styles["Heading3"], fontName=bold_font, fontSize=11, leading=15,
                         spaceBefore=6, spaceAfter=2, textColor=colors.HexColor("#1e3a8a"))
    s_meta = ParagraphStyle("jp_meta", parent=styles["Normal"], fontName=bold_font, fontSize=8,
                            textColor=colors.gray, alignment=1)
    s_box = ParagraphStyle("jp_box", parent=styles["Normal"], fontName=font_name, fontSize=10, leading=16)
    s_list = ParagraphStyle("jp_list", parent=styles["Normal"], fontName=font_name, fontSize=10, leading=16,
                            leftIndent=10)
    s_bq = ParagraphStyle("jp_bq", parent=styles["Normal"], fontName=font_name, fontSize=10, leading=16,
                           textColor=colors.HexColor("#1e40af"))
    s_th = ParagraphStyle("jp_th", parent=styles["Normal"], fontName=bold_font, fontSize=9, leading=13,
                          textColor=colors.white)
    s_td = ParagraphStyle("jp_td", parent=styles["Normal"], fontName=font_name, fontSize=9, leading=13)
    s_td_alt = ParagraphStyle("jp_td_alt", parent=styles["Normal"], fontName=font_name, fontSize=9, leading=13)

    def _xml(text: str) -> str:
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    def _inline(text: str) -> str:
        """Convert **bold** and *italic* to ReportLab XML tags."""
        s = _xml(text)
        s = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', s)
        s = re.sub(r'\*(.+?)\*', r'<i>\1</i>', s)
        return s

    def _colored_box(lines: list[str], bg: str, border: str) -> Table:
        inner_paras = [Paragraph(_inline(ln), s_box) for ln in lines if ln.strip()]
        if not inner_paras:
            inner_paras = [Spacer(1, 1 * mm)]
        inner = [[p] for p in inner_paras]
        content_tbl = Table(inner, colWidths=[page_w - 14])
        content_tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(bg)),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        outer = Table([[content_tbl]], colWidths=[page_w])
        outer.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(bg)),
            ("LINEBEFORE", (0, 0), (0, -1), 4, colors.HexColor(border)),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        return outer

    def _render_table(rows: list[str]) -> Table | None:
        parsed: list[list[str]] = []
        for row in rows:
            if re.match(r'^[\|\s\-:]+$', row):
                continue  # separator row
            cells = [c.strip() for c in row.strip().strip("|").split("|")]
            parsed.append(cells)
        if not parsed:
            return None
        max_cols = max(len(r) for r in parsed)
        col_w = page_w / max_cols
        data = []
        for i, row in enumerate(parsed):
            padded = row + [""] * (max_cols - len(row))
            style = s_th if i == 0 else (s_td_alt if i % 2 == 0 else s_td)
            data.append([Paragraph(_inline(c), style) for c in padded])
        tbl = Table(data, colWidths=[col_w] * max_cols, repeatRows=1)
        ts = [
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e40af")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
        ]
        for ri in range(1, len(data)):
            if ri % 2 == 0:
                ts.append(("BACKGROUND", (0, ri), (-1, ri), colors.HexColor("#f1f5f9")))
        tbl.setStyle(TableStyle(ts))
        return tbl

    # ── block collector ───────────────────────────────────────────────────
    BlockType = str
    Block = dict

    def flush_block(blk: Block, story: list) -> None:
        btype = blk.get("type")
        blines = blk.get("lines", [])
        if not blines and btype not in ("hr", "space"):
            return
        if btype == "space":
            story.append(Spacer(1, 2 * mm))
        elif btype == "hr":
            story.append(HRFlowable(width="100%", thickness=0.8, color=colors.HexColor("#93c5fd"),
                                    dash=(4, 3), spaceAfter=3 * mm, spaceBefore=2 * mm))
        elif btype == "h1":
            story.append(Paragraph(_inline(blines[0]), sh1))
        elif btype == "h2":
            story.append(Paragraph(_inline(blines[0]), sh2))
        elif btype == "h3":
            story.append(Paragraph(_inline(blines[0]), sh3))
        elif btype == "hint":
            story.append(_colored_box(blines, "#fef3c7", "#f59e0b"))
            story.append(Spacer(1, 2 * mm))
        elif btype == "ng":
            story.append(_colored_box(blines, "#fef2f2", "#ef4444"))
            story.append(Spacer(1, 2 * mm))
        elif btype == "ok":
            story.append(_colored_box(blines, "#ecfdf5", "#10b981"))
            story.append(Spacer(1, 2 * mm))
        elif btype == "blockquote":
            story.append(_colored_box(blines, "#eff6ff", "#3b82f6"))
            story.append(Spacer(1, 2 * mm))
        elif btype == "table":
            tbl = _render_table(blines)
            if tbl:
                story.append(tbl)
                story.append(Spacer(1, 2 * mm))
        elif btype == "list":
            for ln in blines:
                story.append(Paragraph(_inline(ln), s_list))
        else:  # para
            for ln in blines:
                if ln.strip():
                    story.append(Paragraph(_inline(ln), sn))

    story: list = []

    # ── page header band (navy gradient simulation) ───────────────────────
    header_style = ParagraphStyle("hp_title", parent=styles["Normal"], fontName=bold_font,
                                  fontSize=14, textColor=colors.white, leading=18)
    header_sub = ParagraphStyle("hp_sub", parent=styles["Normal"], fontName=font_name,
                                fontSize=9, textColor=colors.HexColor("#93c5fd"), leading=13)
    hdr_content = Table(
        [[Paragraph(title, header_style)], [Paragraph(_xml(header), header_sub)]],
        colWidths=[page_w],
    )
    hdr_content.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#0b1e4d")),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
        ("TOPPADDING", (0, 0), (0, 0), 10),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(hdr_content)
    story.append(Spacer(1, 5 * mm))

    # ── parse markdown into blocks ─────────────────────────────────────────
    current: Block = {"type": "para", "lines": []}

    def switch(new_type: BlockType) -> None:
        nonlocal current
        flush_block(current, story)
        current = {"type": new_type, "lines": []}

    for raw_line in markdown_text.splitlines():
        stripped = raw_line.strip()

        if not stripped:
            if current["type"] in ("table", "hint", "ng", "ok", "blockquote"):
                switch("space")
            else:
                if current["lines"]:
                    switch("space")
                else:
                    current = {"type": "space", "lines": []}
                    flush_block(current, story)
                    current = {"type": "para", "lines": []}
            continue

        if stripped.startswith("### "):
            switch("h3")
            current["lines"].append(stripped[4:])
            flush_block(current, story)
            current = {"type": "para", "lines": []}
        elif stripped.startswith("## "):
            switch("h2")
            current["lines"].append(stripped[3:])
            flush_block(current, story)
            current = {"type": "para", "lines": []}
        elif stripped.startswith("# "):
            switch("h1")
            current["lines"].append(stripped[2:])
            flush_block(current, story)
            current = {"type": "para", "lines": []}
        elif stripped.startswith("---") and re.match(r'^-{3,}$', stripped):
            switch("hr")
            flush_block(current, story)
            current = {"type": "para", "lines": []}
        elif stripped.startswith("|"):
            if current["type"] != "table":
                switch("table")
            current["lines"].append(stripped)
        elif stripped.startswith("> ") or stripped == ">":
            if current["type"] != "blockquote":
                switch("blockquote")
            current["lines"].append(stripped[2:] if stripped.startswith("> ") else "")
        elif stripped.startswith("💡"):
            if current["type"] != "hint":
                switch("hint")
            current["lines"].append(stripped)
        elif stripped.startswith("🚫"):
            if current["type"] != "ng":
                switch("ng")
            current["lines"].append(stripped)
        elif stripped.startswith("✅"):
            if current["type"] != "ok":
                switch("ok")
            current["lines"].append(stripped)
        elif re.match(r'^- \[[ x]\]', stripped):
            # checkbox list item
            if current["type"] != "list":
                switch("list")
            checked = stripped[3] == "x"
            rest = stripped[6:]
            sym = "☑" if checked else "☐"
            current["lines"].append(f"{sym} {rest}")
        elif re.match(r'^[-*] ', stripped):
            if current["type"] != "list":
                switch("list")
            current["lines"].append("• " + stripped[2:])
        elif re.match(r'^\d+\. ', stripped):
            if current["type"] != "list":
                switch("list")
            current["lines"].append(stripped)
        else:
            if current["type"] not in ("para",):
                switch("para")
            current["lines"].append(stripped)

    flush_block(current, story)

    # ── page footer ────────────────────────────────────────────────────────
    story.append(Spacer(1, 5 * mm))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cbd5e1")))
    story.append(Paragraph(_xml(footer), s_meta))

    doc.build(story)
    return buf.getvalue()


@app.post("/api/pdf")
async def generate_pdf(body: PdfRequest, request: Request):
    rid = request.state.request_id
    files_map = {
        "01_添削レポート": body.materials.qa,
        "02_自己PR案": body.materials.pr,
        "03_逆質問集": body.materials.questions,
        "04_事前準備チェックリスト": body.materials.checklist,
    }

    present = {k: v for k, v in files_map.items() if v is not None}
    if not present:
        return err("No interview materials provided", status=422, request_id=rid)

    header = f"{body.company_name or '面接対策'} | {body.applicant_name or ''}"
    footer = f"{body.applicant_name or ''} | {body.interview_date or 'TBD'} | {body.job_title or ''}"

    def _content_disposition(filename: str, ext: str) -> str:
        """RFC 5987 encoded Content-Disposition header for non-ASCII filenames."""
        from urllib.parse import quote
        encoded = quote(f"{filename}{ext}", encoding="utf-8")
        ascii_fallback = filename.encode("ascii", errors="replace").decode() + ext
        return f"attachment; filename=\"{ascii_fallback}\"; filename*=UTF-8''{encoded}"

    try:
        # 1ファイルの場合はPDFを直接返す。複数の場合はZIPにまとめる
        if len(present) == 1:
            name, doc = next(iter(present.items()))
            pdf_bytes = _markdown_to_pdf_bytes(name, doc.markdown, header, footer)
            return StreamingResponse(
                BytesIO(pdf_bytes),
                media_type="application/pdf",
                headers={
                    "Content-Disposition": _content_disposition(name, ".pdf"),
                    "X-Request-Id": rid,
                },
            )

        buf = BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            for name, doc in present.items():
                pdf_bytes = _markdown_to_pdf_bytes(name, doc.markdown, header, footer)
                zf.writestr(f"{name}.pdf", pdf_bytes)
        buf.seek(0)

        return StreamingResponse(
            buf,
            media_type="application/zip",
            headers={
                "Content-Disposition": f'attachment; filename="interview_materials.zip"',
                "X-Request-Id": rid,
            },
        )
    except Exception as e:
        import traceback
        tb = traceback.format_exc()
        return err(f"PDF generation failed: {type(e).__name__}: {str(e)[:300]}\nTraceback:\n{tb[-800:]}", status=500, request_id=rid)


# --- /api/interview/pdf  (WeasyPrint スライドPDF生成) ---

_MARP_CSS = """
@page {
  size: 960px 540px;
  margin: 0;
}
* { box-sizing: border-box; }
body { margin: 0; padding: 0; }
.slide-page {
  font-family: 'Hiragino Kaku Gothic ProN', 'Yu Gothic', 'Meiryo', 'Noto Sans JP', sans-serif;
  padding: 32px 46px;
  font-size: 16px;
  line-height: 1.5;
  width: 960px;
  height: 540px;
  overflow: hidden;
  position: relative;
  background: #ffffff;
  color: #1f2937;
  page-break-after: always;
}
.slide-page:last-child { page-break-after: auto; }
.slide-page.title {
  background: linear-gradient(135deg, #0b1e4d 0%, #1e40af 60%, #3b82f6 100%);
  color: #ffffff;
  display: flex;
  flex-direction: column;
  justify-content: center;
  align-items: center;
  text-align: center;
}
.slide-page.title h1 { font-size: 54px; color: #ffffff; border: none; margin-bottom: 16px; }
.slide-page.title h2 { font-size: 24px; color: #bfdbfe; font-weight: normal; border: none; }
.slide-page.title p  { font-size: 20px; color: #dbeafe; margin-top: 20px; }
.slide-page.core {
  background: #0b1e4d;
  color: #ffffff;
  display: flex;
  flex-direction: column;
  justify-content: center;
  align-items: center;
  text-align: center;
}
.slide-page.core h1 { color: #fef3c7; border: none; font-size: 30px; }
.slide-page.core .bigbox {
  background: #1e3a8a;
  border: 3px solid #fbbf24;
  border-radius: 18px;
  padding: 36px 52px;
  margin-top: 24px;
  font-size: 34px;
  font-weight: bold;
  color: #fef3c7;
  max-width: 860px;
  line-height: 1.4;
}
.slide-page.hero {
  background: linear-gradient(135deg, #1e3a8a 0%, #1e40af 100%);
  color: #ffffff;
  display: flex;
  flex-direction: column;
  justify-content: center;
  padding: 32px 46px;
}
.slide-page.hero h1 { color: #fef3c7; border: none; font-size: 28px; }
.slide-page.hero h2 { color: #bfdbfe; font-size: 20px; border: none; border-left: 4px solid #fbbf24; padding-left: 12px; margin-top: 12px; }
.slide-page.hero blockquote { background: rgba(255,255,255,0.1); border-left: 4px solid #fbbf24; color: #f1f5f9; }
.slide-page.hero .star-box { background: rgba(251,191,36,0.2); border: 2px solid #fbbf24; border-radius: 10px; padding: 12px 20px; margin: 8px 0; color: #fef3c7; }
.slide-page.mantra {
  background: #0b1e4d;
  color: #ffffff;
  display: flex;
  flex-direction: column;
  justify-content: center;
  align-items: center;
  text-align: center;
}
.slide-page.mantra h1 { color: #fef3c7; border: none; font-size: 26px; }
.slide-page.mantra .mantra-box {
  background: #1e3a8a;
  border: 2px solid #fbbf24;
  border-radius: 14px;
  padding: 28px 40px;
  margin-top: 20px;
  font-size: 18px;
  color: #f1f5f9;
  line-height: 2.2;
  text-align: left;
  max-width: 860px;
}
.slide-page h1 { color: #0b1e4d; border-bottom: 3px solid #3b82f6; padding-bottom: 6px; font-size: 22px; margin-bottom: 10px; margin-top: 0; }
.slide-page h2 { color: #1e40af; font-size: 18px; border-left: 5px solid #3b82f6; padding-left: 12px; margin-top: 12px; margin-bottom: 4px; }
.slide-page h3 { color: #1e3a8a; font-size: 15px; margin-top: 10px; margin-bottom: 3px; }
.slide-page strong { color: #b91c1c; font-weight: 700; }
.slide-page blockquote { border-left: 5px solid #3b82f6; background: #eff6ff; padding: 10px 16px; margin: 8px 0; border-radius: 0 8px 8px 0; font-size: 15px; line-height: 1.65; color: #1f2937; }
.slide-page table { border-collapse: collapse; margin: 8px 0; width: 100%; font-size: 14px; }
.slide-page th { background: #1e40af; color: #ffffff; padding: 7px 10px; border: 1px solid #1e3a8a; text-align: center; }
.slide-page td { padding: 7px 10px; border: 1px solid #cbd5e1; background: #ffffff; }
.slide-page tr:nth-child(even) td { background: #f1f5f9; }
.slide-page ul, .slide-page ol { font-size: 15px; line-height: 1.7; margin: 4px 0; padding-left: 1.5em; }
.slide-page li { margin-bottom: 3px; }
.slide-page p { margin: 6px 0; font-size: 15px; }
.slide-page .tip { background: #fef3c7; border-left: 5px solid #f59e0b; padding: 7px 14px; margin: 8px 0; border-radius: 0 8px 8px 0; font-size: 13px; color: #7c2d12; }
.slide-page .recommend { background: #ecfdf5; border: 2px solid #10b981; padding: 10px 18px; margin: 8px 0; border-radius: 10px; font-size: 14px; }
.slide-page .ng { background: #fef2f2; border-left: 5px solid #dc2626; padding: 10px 16px; margin: 8px 0; border-radius: 0 8px 8px 0; color: #7f1d1d; font-size: 14px; }
.slide-page .ok { background: #ecfdf5; border-left: 5px solid #10b981; padding: 7px 14px; margin: 8px 0; border-radius: 0 8px 8px 0; font-size: 13px; color: #065f46; }
.slide-page .magnet { background: #fffbeb; border: 3px dashed #f59e0b; padding: 18px 24px; margin: 12px 0; border-radius: 12px; font-size: 18px; color: #78350f; line-height: 2; }
.slide-page .q-card { background: #f8faff; border: 1px solid #bfdbfe; border-radius: 10px; padding: 12px 18px; margin: 8px 0; font-size: 14px; }
.slide-page .q-card h3 { color: #1e40af; margin-top: 0; }
.slide-page .star-row { background: #fffbeb; border: 2px solid #f59e0b; border-radius: 8px; padding: 10px 16px; margin: 6px 0; font-size: 14px; color: #78350f; }
.slide-page .flow { background: #f0f9ff; border: 1px solid #bae6fd; border-radius: 10px; padding: 12px 20px; margin: 8px 0; font-size: 14px; line-height: 1.9; }
.slide-page .sep { border: none; border-top: 2px dashed #93c5fd; margin: 10px 0; }
.slide-page .two-col { display: grid; grid-template-columns: 1fr 1fr; gap: 18px; }
"""


def _marp_md_to_slides_html(markdown_text: str) -> list[tuple[str, str]]:
    """Split Marp markdown by --- and return list of (slide_class, html) tuples."""
    import re
    import markdown2  # type: ignore[import]

    parts = re.split(r'(?:^|\n)---(?:\n|$)', markdown_text)
    slides: list[tuple[str, str]] = []
    for part in parts:
        class_match = re.search(r'<!--\s*_class:\s*(\S+)\s*-->', part)
        slide_class = class_match.group(1) if class_match else ""
        # Strip Marp-specific comments
        cleaned = re.sub(r'<!--[^>]*-->', '', part).strip()
        if not cleaned:
            continue
        html = markdown2.markdown(
            cleaned,
            extras=["fenced-code-blocks", "tables", "break-on-newline", "strike", "task_list"],
        )
        slides.append((slide_class, html))
    return slides


def _slides_to_pdf_bytes(title: str, markdown_text: str) -> bytes:
    """Convert Marp-format markdown to PDF using WeasyPrint."""
    from weasyprint import HTML, CSS  # type: ignore[import]

    slides = _marp_md_to_slides_html(markdown_text)
    if not slides:
        slides = [("", f"<p>{title}</p>")]

    sections = "\n".join(
        f'<section class="slide-page {sc}">{html}</section>'
        for sc, html in slides
    )

    full_html = f"""<!DOCTYPE html>
<html lang="ja">
<head>
<meta charset="utf-8">
</head>
<body>
{sections}
</body>
</html>"""

    pdf_bytes = HTML(string=full_html).write_pdf(
        stylesheets=[CSS(string=_MARP_CSS)]
    )
    return pdf_bytes


@app.post("/api/interview/pdf")
async def interview_slide_pdf(body: InterviewSlidePdfRequest, request: Request):
    rid = request.state.request_id

    doc_map = {
        "01_想定問答集_slides":           body.qa,
        "02_自己PR案_slides":             body.pr,
        "03_逆質問集_slides":             body.questions,
        "04_事前準備チェックリスト_slides": body.checklist,
    }
    present = {k: v for k, v in doc_map.items() if v and v.strip()}
    if not present:
        return err("No slide content provided", status=422, request_id=rid)

    try:
        buf = BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            for name, md in present.items():
                pdf_bytes = _slides_to_pdf_bytes(name, md)
                zf.writestr(f"{name}.pdf", pdf_bytes)
        buf.seek(0)

        return StreamingResponse(
            buf,
            media_type="application/zip",
            headers={
                "Content-Disposition": 'attachment; filename="interview_slides.zip"',
                "X-Request-Id": rid,
            },
        )
    except Exception as e:
        import traceback
        tb = traceback.format_exc()
        return err(
            f"Slide PDF generation failed: {type(e).__name__}: {str(e)[:300]}\n{tb[-600:]}",
            status=500,
            request_id=rid,
        )


# ---------------------------------------------------------------------------
# Entrypoint
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
# Railway rebuild trigger: 2026-05-03 03:07:53 UTC
# Force rebuild at 2026-05-03 03:10:33 UTC
